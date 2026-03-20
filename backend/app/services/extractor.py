"""Textextraktion aus verschiedenen Dateiformaten — PDF, DOCX, XLSX, Bilder (OCR), Text."""

import io
import logging
import time
from dataclasses import dataclass, field

import fitz  # PyMuPDF
import pytesseract
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image

from app.services.exif_stripper import check_for_text_overlays, strip_exif

logger = logging.getLogger(__name__)

# Maximale Dateigrösse: 10 MB
MAX_FILE_SIZE: int = 10 * 1024 * 1024

# Maximale extrahierte Textlaenge: 50.000 Zeichen
MAX_TEXT_LENGTH: int = 50_000

# Unterstuetzte Dateiformate und ihre Zuordnungen
EXTENSION_FORMAT_MAP: dict[str, str] = {
    ".pdf": "PDF",
    ".docx": "DOCX",
    ".xlsx": "XLSX",
    ".txt": "TEXT",
    ".csv": "TEXT",
    ".md": "TEXT",
    ".json": "TEXT",
    ".xml": "TEXT",
    ".html": "TEXT",
    ".png": "IMAGE",
    ".jpg": "IMAGE",
    ".jpeg": "IMAGE",
    ".tiff": "IMAGE",
    ".tif": "IMAGE",
    ".bmp": "IMAGE",
    ".webp": "IMAGE",
}

MIME_FORMAT_MAP: dict[str, str] = {
    "application/pdf": "PDF",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "DOCX",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "XLSX",
    "text/plain": "TEXT",
    "text/csv": "TEXT",
    "text/markdown": "TEXT",
    "text/html": "TEXT",
    "text/xml": "TEXT",
    "application/json": "TEXT",
    "application/xml": "TEXT",
    "image/png": "IMAGE",
    "image/jpeg": "IMAGE",
    "image/tiff": "IMAGE",
    "image/bmp": "IMAGE",
    "image/webp": "IMAGE",
}

# Format-Beschreibungen fuer die API
SUPPORTED_FORMATS: list[dict[str, str]] = [
    {
        "format": "PDF",
        "extensions": ".pdf",
        "description": "PDF-Dokumente — Textextraktion aus allen Seiten",
    },
    {
        "format": "DOCX",
        "extensions": ".docx",
        "description": "Word-Dokumente — Text aus Absaetzen und Tabellen",
    },
    {
        "format": "XLSX",
        "extensions": ".xlsx",
        "description": "Excel-Tabellen — Zellinhalte aller Blaetter",
    },
    {
        "format": "TEXT",
        "extensions": ".txt, .csv, .md, .json, .xml, .html",
        "description": "Textdateien — Direkte UTF-8/Latin-1 Dekodierung",
    },
    {
        "format": "IMAGE",
        "extensions": ".png, .jpg, .jpeg, .tiff, .tif, .bmp, .webp",
        "description": "Bilder — OCR-Texterkennung (Deutsch + Englisch), EXIF-Stripping",
    },
]


@dataclass
class ExtractedText:
    """Ergebnis einer Textextraktion aus einer Datei."""

    text: str
    metadata: dict[str, object] = field(default_factory=dict)
    pages: int = 1
    format: str = ""
    warnings: list[str] = field(default_factory=list)


def _detect_format(filename: str, mime_type: str | None = None) -> str:
    """Erkennt das Dateiformat anhand der Dateiendung (primaer) oder MIME-Type (fallback).

    Args:
        filename: Dateiname mit Endung.
        mime_type: Optionaler MIME-Type.

    Returns:
        Format-String (PDF, DOCX, XLSX, TEXT, IMAGE).

    Raises:
        ValueError: Wenn das Format nicht unterstuetzt wird.
    """
    # Primaer: Dateiendung
    ext = ""
    if "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()

    if ext in EXTENSION_FORMAT_MAP:
        return EXTENSION_FORMAT_MAP[ext]

    # Fallback: MIME-Type
    if mime_type and mime_type in MIME_FORMAT_MAP:
        return MIME_FORMAT_MAP[mime_type]

    supported_extensions = ", ".join(sorted(EXTENSION_FORMAT_MAP.keys()))
    raise ValueError(
        f"Nicht unterstuetztes Dateiformat: '{ext or filename}'. "
        f"Unterstuetzte Formate: {supported_extensions}"
    )


def _extract_pdf(file_bytes: bytes) -> ExtractedText:
    """Extrahiert Text aus einem PDF-Dokument via PyMuPDF.

    Args:
        file_bytes: PDF-Datei als Bytes.

    Returns:
        ExtractedText mit dem extrahierten Text.
    """
    warnings: list[str] = []

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        raise ValueError(f"PDF konnte nicht geoeffnet werden: {e}") from e

    # Verschluesseltes PDF pruefen
    if doc.is_encrypted:
        try:
            # Versuch mit leerem Passwort
            if not doc.authenticate(""):
                doc.close()
                raise ValueError(
                    "Das PDF ist passwortgeschuetzt und kann nicht verarbeitet werden."
                )
            warnings.append("PDF war verschluesselt, konnte aber ohne Passwort geoeffnet werden.")
        except Exception:
            doc.close()
            raise ValueError(
                "Das PDF ist passwortgeschuetzt und kann nicht verarbeitet werden."
            )

    pages: list[str] = []
    page_count = doc.page_count

    for page_num in range(page_count):
        try:
            page = doc[page_num]
            page_text = page.get_text("text")
            if page_text.strip():
                pages.append(page_text.strip())
        except Exception as e:
            warnings.append(f"Seite {page_num + 1} konnte nicht gelesen werden: {e}")

    doc.close()

    text = "\n\n".join(pages)

    if not text.strip():
        warnings.append(
            "Kein Text im PDF gefunden — moeglicherweise handelt es sich um ein gescanntes Dokument."
        )

    metadata: dict[str, object] = {"page_count": page_count}

    return ExtractedText(
        text=text,
        metadata=metadata,
        pages=page_count,
        format="PDF",
        warnings=warnings,
    )


def _extract_docx(file_bytes: bytes) -> ExtractedText:
    """Extrahiert Text aus einem DOCX-Dokument via python-docx.

    Liest Absaetze und Tabelleninhalte aus.

    Args:
        file_bytes: DOCX-Datei als Bytes.

    Returns:
        ExtractedText mit dem extrahierten Text.
    """
    warnings: list[str] = []

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"DOCX konnte nicht geoeffnet werden: {e}") from e

    parts: list[str] = []

    # Absaetze extrahieren
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Tabellen extrahieren
    table_count = 0
    for table in doc.tables:
        table_count += 1
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
        if table_rows:
            parts.append(f"\n[Tabelle {table_count}]\n" + "\n".join(table_rows))

    text = "\n\n".join(parts)

    metadata: dict[str, object] = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": table_count,
    }

    return ExtractedText(
        text=text,
        metadata=metadata,
        pages=1,
        format="DOCX",
        warnings=warnings,
    )


def _extract_xlsx(file_bytes: bytes) -> ExtractedText:
    """Extrahiert Zellinhalte aus einem XLSX-Dokument via openpyxl.

    Verarbeitet alle Blaetter und formatiert den Inhalt als lesbaren Text.

    Args:
        file_bytes: XLSX-Datei als Bytes.

    Returns:
        ExtractedText mit dem extrahierten Text.
    """
    warnings: list[str] = []

    try:
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    except Exception as e:
        raise ValueError(f"XLSX konnte nicht geoeffnet werden: {e}") from e

    parts: list[str] = []
    sheet_count = len(wb.sheetnames)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_lines: list[str] = []

        try:
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                # Nur Zeilen mit mindestens einem nicht-leeren Wert
                if any(c.strip() for c in cells):
                    sheet_lines.append(" | ".join(c for c in cells))
        except Exception as e:
            warnings.append(f"Blatt '{sheet_name}' konnte nicht vollstaendig gelesen werden: {e}")

        if sheet_lines:
            header = f"[Blatt: {sheet_name}]"
            parts.append(header + "\n" + "\n".join(sheet_lines))

    wb.close()

    text = "\n\n".join(parts)

    metadata: dict[str, object] = {
        "sheet_count": sheet_count,
        "sheet_names": wb.sheetnames,
    }

    return ExtractedText(
        text=text,
        metadata=metadata,
        pages=sheet_count,
        format="XLSX",
        warnings=warnings,
    )


def _extract_text_file(file_bytes: bytes) -> ExtractedText:
    """Dekodiert eine Textdatei (UTF-8 mit Latin-1 Fallback).

    Args:
        file_bytes: Textdatei als Bytes.

    Returns:
        ExtractedText mit dem dekodierten Text.
    """
    warnings: list[str] = []
    encoding_used = "utf-8"

    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
            encoding_used = "latin-1"
            warnings.append("Datei wurde mit Latin-1 Encoding gelesen (UTF-8 fehlgeschlagen).")
        except UnicodeDecodeError as e:
            raise ValueError(f"Textdatei konnte nicht dekodiert werden: {e}") from e

    metadata: dict[str, object] = {
        "encoding": encoding_used,
        "line_count": text.count("\n") + 1,
    }

    return ExtractedText(
        text=text,
        metadata=metadata,
        pages=1,
        format="TEXT",
        warnings=warnings,
    )


def _extract_image_ocr(file_bytes: bytes) -> ExtractedText:
    """Extrahiert Text aus einem Bild via OCR (Tesseract) mit EXIF-Stripping.

    Schritt 1: EXIF-Metadaten entfernen (Datenschutz!)
    Schritt 2: Text via Tesseract OCR erkennen (Deutsch + Englisch)

    Args:
        file_bytes: Bilddatei als Bytes.

    Returns:
        ExtractedText mit dem per OCR erkannten Text.
    """
    warnings: list[str] = []

    # Schritt 1: EXIF-Metadaten entfernen
    clean_bytes, exif_metadata = strip_exif(file_bytes)

    if exif_metadata:
        if "gps" in exif_metadata:
            gps = exif_metadata["gps"]
            lat = gps.get("latitude", "?")
            lon = gps.get("longitude", "?")
            warnings.append(
                f"WARNUNG: GPS-Koordinaten in EXIF-Daten gefunden und entfernt "
                f"(Lat: {lat}, Lon: {lon})."
            )
        if "error" not in exif_metadata:
            personal_tags = [
                k for k in exif_metadata
                if k not in ("gps", "error")
            ]
            if personal_tags:
                warnings.append(
                    f"EXIF-Metadaten entfernt: {', '.join(personal_tags[:10])}"
                    + (f" (+{len(personal_tags) - 10} weitere)" if len(personal_tags) > 10 else "")
                )

    # Text-Overlay-Heuristik pruefen
    if check_for_text_overlays(clean_bytes):
        warnings.append(
            "Das Bild enthaelt moeglicherweise Text-Overlays oder Wasserzeichen."
        )

    # Schritt 2: OCR mit Tesseract
    try:
        img = Image.open(io.BytesIO(clean_bytes))
        text = pytesseract.image_to_string(img, lang="deu+eng")
    except pytesseract.TesseractNotFoundError:
        raise ValueError(
            "Tesseract OCR ist nicht installiert. Bild-Texterkennung nicht verfuegbar."
        )
    except Exception as e:
        raise ValueError(f"OCR-Texterkennung fehlgeschlagen: {e}") from e

    text = text.strip()

    if not text:
        warnings.append("Kein Text im Bild erkannt — moeglicherweise ist das Bild zu klein oder undeutlich.")

    metadata: dict[str, object] = {
        "ocr_language": "deu+eng",
        "exif_stripped": bool(exif_metadata and "error" not in exif_metadata),
        "exif_tags_found": len(exif_metadata) if exif_metadata else 0,
    }

    return ExtractedText(
        text=text,
        metadata=metadata,
        pages=1,
        format="IMAGE",
        warnings=warnings,
    )


def extract_text(
    file_bytes: bytes,
    filename: str,
    mime_type: str | None = None,
) -> ExtractedText:
    """Extrahiert Text aus einer Datei basierend auf dem erkannten Format.

    Unterstuetzte Formate: PDF, DOCX, XLSX, TEXT (TXT/CSV/MD/JSON/XML/HTML), IMAGE (OCR).

    Args:
        file_bytes: Dateiinhalt als Bytes.
        filename: Dateiname (fuer Format-Erkennung).
        mime_type: Optionaler MIME-Type (Fallback fuer Format-Erkennung).

    Returns:
        ExtractedText mit Text, Metadaten, Seitenzahl, Format und Warnungen.

    Raises:
        ValueError: Bei nicht unterstuetztem Format, zu grosser Datei
                    oder fehlerhaften Dateien.
    """
    start_time = time.time()

    # Dateigroesse pruefen
    file_size = len(file_bytes)
    if file_size > MAX_FILE_SIZE:
        size_mb = file_size / (1024 * 1024)
        max_mb = MAX_FILE_SIZE / (1024 * 1024)
        raise ValueError(
            f"Datei ist zu gross ({size_mb:.1f} MB). "
            f"Maximale Groesse: {max_mb:.0f} MB."
        )

    if file_size == 0:
        raise ValueError("Die Datei ist leer.")

    # Format erkennen
    detected_format = _detect_format(filename, mime_type)

    # Format-spezifische Extraktion
    extractors = {
        "PDF": _extract_pdf,
        "DOCX": _extract_docx,
        "XLSX": _extract_xlsx,
        "TEXT": _extract_text_file,
        "IMAGE": _extract_image_ocr,
    }

    extractor_fn = extractors[detected_format]
    result = extractor_fn(file_bytes)

    # Textlaenge begrenzen
    if len(result.text) > MAX_TEXT_LENGTH:
        original_length = len(result.text)
        result.text = result.text[:MAX_TEXT_LENGTH]
        result.warnings.append(
            f"Text wurde auf {MAX_TEXT_LENGTH:,} Zeichen gekuerzt "
            f"(Original: {original_length:,} Zeichen)."
        )

    # Metriken loggen
    elapsed = time.time() - start_time
    logger.info(
        "Textextraktion abgeschlossen: format=%s, seiten=%d, zeichen=%d, "
        "dateigroesse=%d bytes, dauer=%.2fs, warnungen=%d",
        result.format,
        result.pages,
        len(result.text),
        file_size,
        elapsed,
        len(result.warnings),
    )

    return result
